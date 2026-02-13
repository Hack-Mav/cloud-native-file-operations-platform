import os
import re
import json
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import logging
import hashlib
from PIL import Image as PILImage
import cv2
import numpy as np
from collections import Counter
import mimetypes

logger = logging.getLogger(__name__)

class ContentAnalyzer:
    """Analyzes and classifies file content using various techniques"""
    
    # File type categories
    FILE_CATEGORIES = {
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
        'video': ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm', '.m4v'],
        'document': ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.odt'],
        'spreadsheet': ['.xlsx', '.xls', '.csv', '.ods'],
        'presentation': ['.pptx', '.ppt', '.odp'],
        'audio': ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a'],
        'archive': ['.zip', '.rar', '.7z', '.tar', '.gz'],
        'code': ['.py', '.js', '.html', '.css', '.java', '.cpp', '.c', '.go', '.rs'],
        'text': ['.txt', '.md', '.log', '.cfg', '.ini', '.json', '.xml', '.yaml', '.yml']
    }
    
    # Content classification patterns
    CONTENT_PATTERNS = {
        'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        'phone': r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b',
        'url': r'https?://(?:[-\w.])+(?:[:\d]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:\w*))?)?',
        'credit_card': r'\b(?:\d{4}[-\s]?){3}\d{4}\b',
        'ssn': r'\b\d{3}-\d{2}-\d{4}\b',
        'ip_address': r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b',
        'date': r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',
        'currency': r'\$\d+(?:,\d{3})*(?:\.\d{2})?|\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:USD|EUR|GBP|JPY)',
    }
    
    # Sensitive content patterns
    SENSITIVE_PATTERNS = {
        'api_key': r'(?i)(api[_-]?key|apikey)\s*[:=]\s*["\']?[a-zA-Z0-9_-]{16,}["\']?',
        'password': r'(?i)(password|passwd|pwd)\s*[:=]\s*["\']?[^\s"\']{6,}["\']?',
        'token': r'(?i)(token|access[_-]?token)\s*[:=]\s*["\']?[a-zA-Z0-9._-]{20,}["\']?',
        'secret': r'(?i)(secret|private[_-]?key)\s*[:=]\s*["\']?[a-zA-Z0-9._-]{16,}["\']?',
    }
    
    def __init__(self, temp_dir: str = "/tmp/processing"):
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    async def analyze_file(
        self,
        file_path: str,
        extract_metadata: bool = True,
        scan_for_sensitive: bool = True,
        content_classification: bool = True
    ) -> Dict[str, Any]:
        """
        Perform comprehensive file analysis
        
        Args:
            file_path: Path to file to analyze
            extract_metadata: Whether to extract file metadata
            scan_for_sensitive: Whether to scan for sensitive content
            content_classification: Whether to classify content type
        
        Returns:
            Dict with analysis results
        """
        try:
            results = {
                'success': True,
                'file_path': file_path,
                'file_info': {},
                'content_analysis': {},
                'security_analysis': {},
                'classification': {}
            }
            
            # Basic file information
            file_info = await self._get_file_info(file_path)
            results['file_info'] = file_info
            
            # Content analysis based on file type
            category = file_info.get('category')
            
            if category == 'image':
                content_analysis = await self._analyze_image(file_path)
            elif category == 'video':
                content_analysis = await self._analyze_video(file_path)
            elif category == 'document':
                content_analysis = await self._analyze_document(file_path)
            elif category == 'text':
                content_analysis = await self._analyze_text_file(file_path)
            else:
                content_analysis = {'message': f'No specific analysis for category: {category}'}
            
            results['content_analysis'] = content_analysis
            
            # Security analysis
            if scan_for_sensitive:
                security_analysis = await self._security_scan(file_path, category)
                results['security_analysis'] = security_analysis
            
            # Content classification
            if content_classification:
                classification = await self._classify_content(file_path, category, content_analysis)
                results['classification'] = classification
            
            return results
            
        except Exception as e:
            logger.error(f"Error analyzing file {file_path}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        try:
            file_path = Path(file_path)
            stat = file_path.stat()
            
            # Determine file category
            ext = file_path.suffix.lower()
            category = 'unknown'
            for cat, extensions in self.FILE_CATEGORIES.items():
                if ext in extensions:
                    category = cat
                    break
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            # Calculate file hashes
            file_hash = await self._calculate_file_hash(str(file_path))
            
            return {
                'name': file_path.name,
                'size_bytes': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'extension': ext,
                'category': category,
                'mime_type': mime_type,
                'created': stat.st_ctime,
                'modified': stat.st_mtime,
                'md5_hash': file_hash.get('md5'),
                'sha256_hash': file_hash.get('sha256')
            }
            
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {'error': str(e)}
    
    async def _calculate_file_hash(self, file_path: str) -> Dict[str, str]:
        """Calculate MD5 and SHA256 hashes"""
        try:
            md5_hash = hashlib.md5()
            sha256_hash = hashlib.sha256()
            
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    md5_hash.update(chunk)
                    sha256_hash.update(chunk)
            
            return {
                'md5': md5_hash.hexdigest(),
                'sha256': sha256_hash.hexdigest()
            }
            
        except Exception as e:
            logger.error(f"Error calculating file hash: {str(e)}")
            return {}
    
    async def _analyze_image(self, image_path: str) -> Dict[str, Any]:
        """Analyze image content"""
        try:
            with PILImage.open(image_path) as img:
                # Basic image properties
                properties = {
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info,
                    'color_palette': len(img.getcolors(maxcolors=256*256*256)) if img.mode != 'P' else len(img.getpalette() or [])
                }
                
                # Advanced analysis using OpenCV
                try:
                    # Convert to OpenCV format
                    img_cv = cv2.imread(image_path)
                    if img_cv is not None:
                        # Color analysis
                        avg_color = np.mean(img_cv, axis=(0, 1))
                        properties['average_color'] = {
                            'blue': float(avg_color[0]),
                            'green': float(avg_color[1]),
                            'red': float(avg_color[2])
                        }
                        
                        # Brightness analysis
                        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                        properties['brightness'] = {
                            'mean': float(np.mean(gray)),
                            'std': float(np.std(gray)),
                            'min': float(np.min(gray)),
                            'max': float(np.max(gray))
                        }
                        
                        # Edge detection for content complexity
                        edges = cv2.Canny(gray, 50, 150)
                        edge_density = np.sum(edges > 0) / edges.size
                        properties['edge_density'] = float(edge_density)
                        
                        # Detect if image is likely a screenshot
                        if edge_density > 0.1:
                            properties['likely_screenshot'] = True
                        
                        # Detect dominant colors using K-means
                        try:
                            pixels = img_cv.reshape(-1, 3)
                            from sklearn.cluster import KMeans
                            kmeans = KMeans(n_clusters=5, random_state=42, n_init=10)
                            kmeans.fit(pixels)
                            dominant_colors = kmeans.cluster_centers_.astype(int)
                            properties['dominant_colors'] = [
                                {'r': int(color[2]), 'g': int(color[1]), 'b': int(color[0])}
                                for color in dominant_colors
                            ]
                        except ImportError:
                            # Fallback if sklearn not available
                            properties['dominant_colors'] = []
                
                except Exception as e:
                    logger.warning(f"OpenCV analysis failed: {str(e)}")
                
                # EXIF data extraction
                exif_data = {}
                if hasattr(img, '_getexif') and img._getexif():
                    try:
                        from PIL.ExifTags import TAGS
                        exif = img._getexif()
                        for tag_id, value in exif.items():
                            tag = TAGS.get(tag_id, tag_id)
                            exif_data[tag] = value
                    except:
                        pass
                
                properties['exif_data'] = exif_data
                
                return {
                    'type': 'image',
                    'properties': properties
                }
                
        except Exception as e:
            logger.error(f"Error analyzing image {image_path}: {str(e)}")
            return {'error': str(e)}
    
    async def _analyze_video(self, video_path: str) -> Dict[str, Any]:
        """Analyze video content"""
        try:
            cap = cv2.VideoCapture(video_path)
            
            if not cap.isOpened():
                return {'error': 'Could not open video file'}
            
            # Get video properties
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0
            
            properties = {
                'fps': fps,
                'frame_count': frame_count,
                'resolution': (width, height),
                'duration_seconds': duration,
                'duration_formatted': self._format_duration(duration)
            }
            
            # Sample frames for content analysis
            sample_frames = []
            sample_interval = max(1, frame_count // 10)  # Sample 10 frames
            
            for i in range(0, frame_count, sample_interval):
                cap.set(cv2.CAP_PROP_POS_FRAMES, i)
                ret, frame = cap.read()
                if ret:
                    # Analyze frame
                    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                    brightness = np.mean(gray)
                    edge_density = np.sum(cv2.Canny(gray, 50, 150) > 0) / gray.size
                    
                    sample_frames.append({
                        'frame_number': i,
                        'timestamp': i / fps,
                        'brightness': float(brightness),
                        'edge_density': float(edge_density)
                    })
                    
                    if len(sample_frames) >= 10:
                        break
            
            cap.release()
            
            # Analyze frame samples
            if sample_frames:
                avg_brightness = np.mean([f['brightness'] for f in sample_frames])
                avg_edge_density = np.mean([f['edge_density'] for f in sample_frames])
                
                properties.update({
                    'average_brightness': float(avg_brightness),
                    'average_edge_density': float(avg_edge_density),
                    'likely_animation': avg_edge_density > 0.15,
                    'samples_analyzed': len(sample_frames)
                })
            
            return {
                'type': 'video',
                'properties': properties,
                'frame_samples': sample_frames
            }
            
        except Exception as e:
            logger.error(f"Error analyzing video {video_path}: {str(e)}")
            return {'error': str(e)}
    
    async def _analyze_document(self, document_path: str) -> Dict[str, Any]:
        """Analyze document content"""
        try:
            ext = Path(document_path).suffix.lower()
            
            if ext == '.pdf':
                return await self._analyze_pdf(document_path)
            elif ext in ['.docx', '.doc']:
                return await self._analyze_docx(document_path)
            elif ext in ['.txt', '.md']:
                return await self._analyze_text_file(document_path)
            else:
                return {'message': f'Document analysis not implemented for {ext}'}
                
        except Exception as e:
            logger.error(f"Error analyzing document {document_path}: {str(e)}")
            return {'error': str(e)}
    
    async def _analyze_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Analyze PDF content"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(pdf_path)
            
            # Extract text content
            full_text = ""
            page_texts = []
            
            for page_num in range(min(doc.page_count, 20)):  # Limit to first 20 pages
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    page_texts.append({
                        'page': page_num + 1,
                        'text': text,
                        'word_count': len(text.split()),
                        'char_count': len(text)
                    })
                    full_text += text + "\n"
            
            # Analyze text content
            text_analysis = await self._analyze_text_content(full_text)
            
            # Document metadata
            metadata = {
                'page_count': doc.page_count,
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
                'modification_date': doc.metadata.get('modDate', '')
            }
            
            doc.close()
            
            return {
                'type': 'pdf',
                'metadata': metadata,
                'text_analysis': text_analysis,
                'page_analysis': page_texts,
                'total_words': sum(page['word_count'] for page in page_texts),
                'total_characters': sum(page['char_count'] for page in page_texts)
            }
            
        except ImportError:
            return {'error': 'PyMuPDF not available for PDF analysis'}
        except Exception as e:
            logger.error(f"Error analyzing PDF {pdf_path}: {str(e)}")
            return {'error': str(e)}
    
    async def _analyze_docx(self, docx_path: str) -> Dict[str, Any]:
        """Analyze DOCX content"""
        try:
            import docx
            
            doc = docx.Document(docx_path)
            
            # Extract text
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'word_count': len(para.text.split())
                    })
                    full_text += para.text + "\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Analyze text content
            text_analysis = await self._analyze_text_content(full_text)
            
            # Document metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or ''
            }
            
            return {
                'type': 'docx',
                'metadata': metadata,
                'text_analysis': text_analysis,
                'paragraphs': paragraphs,
                'tables': tables,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables)
            }
            
        except ImportError:
            return {'error': 'python-docx not available for DOCX analysis'}
        except Exception as e:
            logger.error(f"Error analyzing DOCX {docx_path}: {str(e)}")
            return {'error': str(e)}
    
    async def _analyze_text_file(self, text_path: str) -> Dict[str, Any]:
        """Analyze text file content"""
        try:
            with open(text_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            return await self._analyze_text_content(content)
            
        except Exception as e:
            logger.error(f"Error analyzing text file {text_path}: {str(e)}")
            return {'error': str(e)}
    
    async def _analyze_text_content(self, content: str) -> Dict[str, Any]:
        """Analyze text content"""
        try:
            # Basic statistics
            words = content.split()
            sentences = re.split(r'[.!?]+', content)
            lines = content.split('\n')
            
            analysis = {
                'character_count': len(content),
                'word_count': len(words),
                'sentence_count': len([s for s in sentences if s.strip()]),
                'line_count': len(lines),
                'paragraph_count': len([p for p in content.split('\n\n') if p.strip()]),
                'average_word_length': np.mean([len(word) for word in words]) if words else 0,
                'average_sentence_length': np.mean([len(s.split()) for s in sentences if s.strip()]) if sentences else 0
            }
            
            # Language detection (simple)
            language = self._detect_language(content)
            analysis['detected_language'] = language
            
            # Pattern matching
            patterns_found = {}
            for pattern_name, pattern in self.CONTENT_PATTERNS.items():
                matches = re.findall(pattern, content)
                if matches:
                    patterns_found[pattern_name] = {
                        'count': len(matches),
                        'samples': matches[:5]  # First 5 samples
                    }
            
            analysis['patterns_found'] = patterns_found
            
            # Keyword extraction
            keywords = self._extract_keywords(content)
            analysis['keywords'] = keywords
            
            # Readability score (simplified)
            readability = self._calculate_readability(content)
            analysis['readability'] = readability
            
            return {
                'type': 'text',
                'analysis': analysis
            }
            
        except Exception as e:
            logger.error(f"Error analyzing text content: {str(e)}")
            return {'error': str(e)}
    
    async def _security_scan(self, file_path: str, category: str) -> Dict[str, Any]:
        """Scan file for sensitive content"""
        try:
            security_analysis = {
                'sensitive_patterns': {},
                'risk_level': 'low',
                'recommendations': []
            }
            
            # For text-based files, scan content
            if category in ['text', 'document', 'code']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    
                    # Scan for sensitive patterns
                    for pattern_name, pattern in self.SENSITIVE_PATTERNS.items():
                        matches = re.findall(pattern, content, re.IGNORECASE)
                        if matches:
                            security_analysis['sensitive_patterns'][pattern_name] = {
                                'count': len(matches),
                                'samples': matches[:3]
                            }
                    
                    # Determine risk level
                    total_sensitive = sum(len(data['samples']) for data in security_analysis['sensitive_patterns'].values())
                    if total_sensitive > 10:
                        security_analysis['risk_level'] = 'high'
                    elif total_sensitive > 3:
                        security_analysis['risk_level'] = 'medium'
                    
                    # Add recommendations
                    if security_analysis['sensitive_patterns']:
                        security_analysis['recommendations'].append('Review and remove sensitive information')
                        security_analysis['recommendations'].append('Consider using environment variables for secrets')
                
                except Exception as e:
                    logger.warning(f"Could not scan file content: {str(e)}")
            
            # Check file name for sensitive information
            filename = Path(file_path).name.lower()
            sensitive_keywords = ['password', 'secret', 'key', 'token', 'private', 'confidential']
            if any(keyword in filename for keyword in sensitive_keywords):
                security_analysis['risk_level'] = 'medium'
                security_analysis['recommendations'].append('Consider renaming file to avoid exposing sensitive information')
            
            return security_analysis
            
        except Exception as e:
            logger.error(f"Error in security scan: {str(e)}")
            return {'error': str(e)}
    
    async def _classify_content(self, file_path: str, category: str, content_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Classify content into specific categories"""
        try:
            classification = {
                'primary_category': category,
                'subcategories': [],
                'tags': [],
                'confidence': 0.8
            }
            
            # Image classification
            if category == 'image' and 'properties' in content_analysis:
                props = content_analysis['properties']
                
                # Classify by properties
                if props.get('edge_density', 0) > 0.15:
                    classification['subcategories'].append('complex')
                else:
                    classification['subcategories'].append('simple')
                
                if props.get('has_transparency'):
                    classification['subcategories'].append('transparent')
                
                if props.get('mode') == 'RGB':
                    classification['subcategories'].append('color')
                elif props.get('mode') == 'L':
                    classification['subcategories'].append('grayscale')
                
                # Size-based classification
                size = props.get('size', (0, 0))
                if size[0] > 1920 or size[1] > 1080:
                    classification['subcategories'].append('high_resolution')
                elif size[0] < 500 or size[1] < 500:
                    classification['subcategories'].append('thumbnail')
            
            # Text classification
            elif category in ['text', 'document'] and 'analysis' in content_analysis:
                analysis = content_analysis['analysis']
                
                # Classify by content type
                patterns = analysis.get('patterns_found', {})
                if patterns.get('email'):
                    classification['subcategories'].append('contains_emails')
                if patterns.get('phone'):
                    classification['subcategories'].append('contains_phones')
                if patterns.get('url'):
                    classification['subcategories'].append('contains_urls')
                if patterns.get('credit_card'):
                    classification['subcategories'].append('contains_financial')
                
                # Classify by length
                word_count = analysis.get('word_count', 0)
                if word_count > 10000:
                    classification['subcategories'].append('long_document')
                elif word_count < 100:
                    classification['subcategories'].append('short_text')
                
                # Language-based tags
                language = analysis.get('detected_language', 'unknown')
                if language != 'unknown':
                    classification['tags'].append(f'language:{language}')
            
            # Video classification
            elif category == 'video' and 'properties' in content_analysis:
                props = content_analysis['properties']
                
                duration = props.get('duration_seconds', 0)
                if duration > 600:  # 10 minutes
                    classification['subcategories'].append('long_video')
                elif duration < 60:  # 1 minute
                    classification['subcategories'].append('short_clip')
                
                if props.get('likely_animation'):
                    classification['subcategories'].append('animation')
                
                resolution = props.get('resolution', (0, 0))
                if resolution[0] >= 1920 and resolution[1] >= 1080:
                    classification['subcategories'].append('hd')
                elif resolution[0] >= 3840 and resolution[1] >= 2160:
                    classification['subcategories'].append('4k')
            
            # Add general tags
            filename = Path(file_path).name.lower()
            if 'test' in filename:
                classification['tags'].append('test_file')
            if 'sample' in filename:
                classification['tags'].append('sample')
            if 'backup' in filename:
                classification['tags'].append('backup')
            
            return classification
            
        except Exception as e:
            logger.error(f"Error in content classification: {str(e)}")
            return {'error': str(e)}
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection based on character patterns"""
        try:
            # Simple heuristic based on common characters
            if re.search(r'[\u4e00-\u9fff]', text):  # Chinese characters
                return 'chinese'
            elif re.search(r'[а-яё]', text, re.IGNORECASE):  # Cyrillic
                return 'russian'
            elif re.search(r'[ñáéíóúü]', text, re.IGNORECASE):  # Spanish accents
                return 'spanish'
            elif re.search(r'[äöüß]', text, re.IGNORECASE):  # German characters
                return 'german'
            elif re.search(r'[àâæçéèêëîïôœùûüÿ]', text, re.IGNORECASE):  # French accents
                return 'french'
            else:
                return 'english'  # Default assumption
        except:
            return 'unknown'
    
    def _extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from text"""
        try:
            # Simple keyword extraction based on word frequency
            words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
            
            # Filter out common stop words
            stop_words = {'the', 'and', 'that', 'have', 'for', 'not', 'with', 'you', 'this', 'but', 'his', 'from', 'they', 'she', 'her', 'been', 'than', 'its', 'were', 'said', 'each', 'which', 'their', 'time', 'will', 'about', 'would', 'there', 'could', 'other', 'after', 'first', 'think', 'more', 'very', 'what', 'when', 'make', 'like', 'can', 'just', 'know', 'take', 'people', 'year', 'your', 'good', 'some', 'could', 'them', 'see', 'other', 'than', 'then', 'now', 'look', 'only', 'come', 'its', 'over', 'think', 'also', 'back', 'after', 'use', 'two', 'how', 'our', 'work', 'first', 'well', 'way', 'even', 'new', 'want', 'because', 'any', 'these', 'give', 'day', 'most', 'us', 'is', 'was', 'are', 'been', 'has', 'had', 'were', 'said', 'did', 'having', 'may', 'am'}
            
            filtered_words = [word for word in words if word not in stop_words and len(word) > 3]
            
            # Count word frequencies
            word_freq = Counter(filtered_words)
            
            # Return most common keywords
            return [word for word, count in word_freq.most_common(max_keywords)]
            
        except Exception as e:
            logger.error(f"Error extracting keywords: {str(e)}")
            return []
    
    def _calculate_readability(self, text: str) -> Dict[str, Any]:
        """Calculate basic readability metrics"""
        try:
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            if not sentences:
                return {'score': 0, 'level': 'unknown'}
            
            words = text.split()
            avg_sentence_length = len(words) / len(sentences)
            
            # Simplified readability score (lower is easier to read)
            score = avg_sentence_length + (sum(len(word) for word in words) / len(words)) * 0.5
            
            if score < 10:
                level = 'very_easy'
            elif score < 15:
                level = 'easy'
            elif score < 20:
                level = 'medium'
            elif score < 25:
                level = 'difficult'
            else:
                level = 'very_difficult'
            
            return {
                'score': round(score, 2),
                'level': level,
                'avg_sentence_length': round(avg_sentence_length, 2),
                'avg_word_length': round(sum(len(word) for word in words) / len(words), 2) if words else 0
            }
            
        except Exception as e:
            logger.error(f"Error calculating readability: {str(e)}")
            return {'error': str(e)}
    
    def _format_duration(self, seconds: float) -> str:
        """Format duration in seconds to human-readable format"""
        try:
            hours = int(seconds // 3600)
            minutes = int((seconds % 3600) // 60)
            secs = int(seconds % 60)
            
            if hours > 0:
                return f"{hours:02d}:{minutes:02d}:{secs:02d}"
            else:
                return f"{minutes:02d}:{secs:02d}"
        except:
            return "00:00"
