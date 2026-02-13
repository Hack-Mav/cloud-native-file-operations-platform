import os
import tempfile
import io
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import logging
import fitz  # PyMuPDF
import docx
import openpyxl
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import pdfplumber
import pandas as pd
from PIL import Image as PILImage

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing including text extraction and PDF generation"""
    
    SUPPORTED_FORMATS = {
        'input': ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.xlsx', '.xls', '.csv', '.pptx', '.ppt'],
        'output': ['.pdf', '.txt', '.docx', '.html']
    }
    
    def __init__(self, temp_dir: str = "/tmp/processing"):
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    async def extract_text_from_pdf(
        self, 
        pdf_path: str,
        extract_images: bool = False,
        preserve_layout: bool = True
    ) -> Dict[str, Any]:
        """
        Extract text and optionally images from PDF document
        
        Args:
            pdf_path: Path to PDF file
            extract_images: Whether to extract images
            preserve_layout: Whether to preserve original layout
        
        Returns:
            Dict with extracted content and metadata
        """
        try:
            text_content = []
            images = []
            metadata = {}
            
            # Extract using PyMuPDF for better layout preservation
            if preserve_layout:
                doc = fitz.open(pdf_path)
                metadata.update({
                    'page_count': doc.page_count,
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', ''),
                    'creator': doc.metadata.get('creator', ''),
                    'producer': doc.metadata.get('producer', ''),
                    'creation_date': doc.metadata.get('creationDate', ''),
                    'modification_date': doc.metadata.get('modDate', '')
                })
                
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    
                    # Extract text with layout
                    text = page.get_text()
                    if text.strip():
                        text_content.append({
                            'page': page_num + 1,
                            'text': text,
                            'bbox': page.rect
                        })
                    
                    # Extract images if requested
                    if extract_images:
                        image_list = page.get_images()
                        for img_index, img in enumerate(image_list):
                            try:
                                xref = img[0]
                                pix = fitz.Pixmap(doc, xref)
                                
                                if pix.n - pix.alpha < 4:  # GRAY or RGB
                                    img_data = pix.tobytes("png")
                                    img_filename = f"page_{page_num + 1}_img_{img_index + 1}.png"
                                    img_path = self.temp_dir / img_filename
                                    
                                    with open(img_path, "wb") as f:
                                        f.write(img_data)
                                    
                                    images.append({
                                        'page': page_num + 1,
                                        'filename': img_filename,
                                        'path': str(img_path),
                                        'size_bytes': len(img_data),
                                        'width': pix.width,
                                        'height': pix.height
                                    })
                                
                                pix = None
                            except Exception as e:
                                logger.warning(f"Failed to extract image from page {page_num + 1}: {str(e)}")
                
                doc.close()
            
            # Alternative extraction using pdfplumber for tables
            tables = []
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table_index, table in enumerate(page_tables):
                                tables.append({
                                    'page': page_num + 1,
                                    'table_index': table_index,
                                    'data': table,
                                    'rows': len(table),
                                    'columns': len(table[0]) if table else 0
                                })
            except Exception as e:
                logger.warning(f"Failed to extract tables using pdfplumber: {str(e)}")
            
            return {
                'success': True,
                'text_content': text_content,
                'images': images,
                'tables': tables,
                'metadata': metadata,
                'total_pages': len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def extract_text_from_docx(
        self, 
        docx_path: str,
        extract_images: bool = False
    ) -> Dict[str, Any]:
        """
        Extract text and optionally images from DOCX document
        
        Args:
            docx_path: Path to DOCX file
            extract_images: Whether to extract images
        
        Returns:
            Dict with extracted content and metadata
        """
        try:
            doc = docx.Document(docx_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract tables
            tables = []
            for table_index, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                tables.append({
                    'table_index': table_index,
                    'data': table_data,
                    'rows': len(table_data),
                    'columns': len(table_data[0]) if table_data else 0
                })
            
            # Extract images if requested
            images = []
            if extract_images:
                try:
                    # Extract images from docx
                    import zipfile
                    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                        image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                        
                        for img_index, img_file in enumerate(image_files):
                            img_data = docx_zip.read(img_file)
                            img_filename = f"docx_img_{img_index + 1}.{Path(img_file).suffix}"
                            img_path = self.temp_dir / img_filename
                            
                            with open(img_path, "wb") as f:
                                f.write(img_data)
                            
                            images.append({
                                'filename': img_filename,
                                'path': str(img_path),
                                'size_bytes': len(img_data),
                                'original_path': img_file
                            })
                except Exception as e:
                    logger.warning(f"Failed to extract images from DOCX: {str(e)}")
            
            # Document metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or ''
            }
            
            return {
                'success': True,
                'paragraphs': paragraphs,
                'tables': tables,
                'images': images,
                'metadata': metadata,
                'total_paragraphs': len(paragraphs)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def extract_text_from_excel(
        self, 
        excel_path: str,
        sheet_names: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Extract data from Excel files
        
        Args:
            excel_path: Path to Excel file
            sheet_names: Specific sheets to extract (None for all)
        
        Returns:
            Dict with extracted data and metadata
        """
        try:
            # Try openpyxl first for .xlsx files
            if excel_path.endswith('.xlsx'):
                workbook = openpyxl.load_workbook(excel_path, data_only=True)
                
                sheets_data = {}
                metadata = {
                    'sheet_names': workbook.sheetnames,
                    'total_sheets': len(workbook.sheetnames)
                }
                
                for sheet_name in workbook.sheetnames:
                    if sheet_names and sheet_name not in sheet_names:
                        continue
                    
                    sheet = workbook[sheet_name]
                    data = []
                    
                    for row in sheet.iter_rows(values_only=True):
                        # Filter out None values at the end of rows
                        if row and any(cell is not None for cell in row):
                            data.append(list(row))
                    
                    sheets_data[sheet_name] = {
                        'data': data,
                        'rows': len(data),
                        'columns': len(data[0]) if data else 0,
                        'max_row': sheet.max_row,
                        'max_column': sheet.max_column
                    }
                
                workbook.close()
                
                return {
                    'success': True,
                    'sheets': sheets_data,
                    'metadata': metadata
                }
            
            # Use pandas for CSV files
            elif excel_path.endswith('.csv'):
                df = pd.read_csv(excel_path)
                
                return {
                    'success': True,
                    'data': df.to_dict('records'),
                    'columns': df.columns.tolist(),
                    'rows': len(df),
                    'metadata': {
                        'file_type': 'CSV',
                        'encoding': 'utf-8'  # Default, could be detected
                    }
                }
            
            else:
                return {
                    'success': False,
                    'error': f'Unsupported Excel format: {Path(excel_path).suffix}'
                }
                
        except Exception as e:
            logger.error(f"Error extracting data from Excel {excel_path}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def generate_pdf_from_text(
        self,
        text_content: Union[str, List[str]],
        output_path: str,
        title: str = "Generated Document",
        author: str = "",
        font_size: int = 12,
        page_size: str = "A4"
    ) -> Dict[str, Any]:
        """
        Generate PDF from text content
        
        Args:
            text_content: Text content (string or list of paragraphs)
            output_path: Path for output PDF
            title: Document title
            author: Document author
            font_size: Base font size
            page_size: Page size (A4, letter)
        
        Returns:
            Dict with generation results
        """
        try:
            # Setup page size
            if page_size.upper() == "A4":
                pagesize = A4
            else:
                pagesize = letter
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=pagesize,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            normal_style = styles['Normal']
            
            # Create custom styles
            title_style.fontSize = font_size + 8
            title_style.leading = font_size + 12
            normal_style.fontSize = font_size
            normal_style.leading = font_size + 4
            
            # Build content
            story = []
            
            # Add title
            if title:
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 12))
            
            # Add text content
            if isinstance(text_content, str):
                paragraphs = text_content.split('\n\n')
            else:
                paragraphs = text_content
            
            for para in paragraphs:
                if para.strip():
                    # Handle long paragraphs by splitting them
                    words = para.split()
                    current_para = ""
                    
                    for word in words:
                        test_para = current_para + " " + word if current_para else word
                        # Simple length check (could be improved with actual width calculation)
                        if len(test_para) < 1000:  # Rough character limit
                            current_para = test_para
                        else:
                            if current_para:
                                story.append(Paragraph(current_para, normal_style))
                                story.append(Spacer(1, 6))
                            current_para = word
                    
                    if current_para:
                        story.append(Paragraph(current_para, normal_style))
                        story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            # Get file size
            file_size = os.path.getsize(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'file_size_bytes': file_size,
                'pages': len(story),
                'title': title,
                'author': author,
                'font_size': font_size,
                'page_size': page_size
            }
            
        except Exception as e:
            logger.error(f"Error generating PDF from text: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def generate_pdf_from_images(
        self,
        image_paths: List[str],
        output_path: str,
        title: str = "Image Document",
        fit_to_page: bool = True,
        image_quality: float = 0.8
    ) -> Dict[str, Any]:
        """
        Generate PDF from image files
        
        Args:
            image_paths: List of image file paths
            output_path: Path for output PDF
            title: Document title
            fit_to_page: Whether to fit images to page
            image_quality: Image quality for compression
        
        Returns:
            Dict with generation results
        """
        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            
            # Add title page
            if title:
                styles = getSampleStyleSheet()
                title_style = styles['Title']
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 50))
            
            # Process each image
            for img_path in image_paths:
                try:
                    # Open and process image
                    with PILImage.open(img_path) as img:
                        # Convert to RGB if necessary
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Create temporary image file for PDF
                        temp_img_path = self.temp_dir / f"temp_{Path(img_path).stem}.jpg"
                        
                        # Resize if needed
                        if fit_to_page:
                            img.thumbnail((7 * inch, 9 * inch), PILImage.Resampling.LANCZOS)
                        
                        img.save(temp_img_path, 'JPEG', quality=int(image_quality * 100))
                        
                        # Add to story
                        rl_img = RLImage(str(temp_img_path), width=6*inch, height=8*inch)
                        story.append(rl_img)
                        story.append(Spacer(1, 20))
                        
                        # Clean up temp file
                        temp_img_path.unlink()
                
                except Exception as e:
                    logger.warning(f"Failed to process image {img_path}: {str(e)}")
                    continue
            
            # Build PDF
            doc.build(story)
            
            file_size = os.path.getsize(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'file_size_bytes': file_size,
                'images_processed': len(image_paths),
                'title': title
            }
            
        except Exception as e:
            logger.error(f"Error generating PDF from images: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def convert_document_format(
        self,
        input_path: str,
        output_path: str,
        target_format: str
    ) -> Dict[str, Any]:
        """
        Convert document to different format
        
        Args:
            input_path: Path to input document
            output_path: Path for output document
            target_format: Target format (pdf, txt, html)
        
        Returns:
            Dict with conversion results
        """
        try:
            input_ext = Path(input_path).suffix.lower()
            target_ext = target_format.lower()
            
            # PDF to Text
            if input_ext == '.pdf' and target_ext in ['.txt', '.text']:
                result = await self.extract_text_from_pdf(input_path)
                if result['success']:
                    # Combine all text
                    full_text = ""
                    for page_content in result['text_content']:
                        full_text += page_content['text'] + "\n\n"
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(full_text)
                    
                    return {
                        'success': True,
                        'output_path': output_path,
                        'original_format': 'PDF',
                        'target_format': target_format.upper(),
                        'characters_extracted': len(full_text)
                    }
            
            # DOCX to PDF
            elif input_ext == '.docx' and target_ext == '.pdf':
                result = await self.extract_text_from_docx(input_path)
                if result['success']:
                    # Combine paragraphs
                    paragraphs = [para['text'] for para in result['paragraphs']]
                    title = result['metadata'].get('title', 'Converted Document')
                    author = result['metadata'].get('author', '')
                    
                    return await self.generate_pdf_from_text(
                        paragraphs, output_path, title, author
                    )
            
            # Text to PDF
            elif input_ext in ['.txt', '.text'] and target_ext == '.pdf':
                with open(input_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                
                return await self.generate_pdf_from_text(
                    text_content, output_path, Path(input_path).stem
                )
            
            # Excel to PDF (as table data)
            elif input_ext in ['.xlsx', '.xls'] and target_ext == '.pdf':
                result = await self.extract_text_from_excel(input_path)
                if result['success']:
                    # Convert each sheet to text
                    all_text = ""
                    for sheet_name, sheet_data in result.get('sheets', {}).items():
                        all_text += f"Sheet: {sheet_name}\n\n"
                        for row in sheet_data['data']:
                            all_text += "\t".join(str(cell) if cell is not None else "" for cell in row) + "\n"
                        all_text += "\n"
                    
                    return await self.generate_pdf_from_text(
                        all_text, output_path, f"Excel Document - {Path(input_path).stem}"
                    )
            
            else:
                return {
                    'success': False,
                    'error': f'Conversion from {input_ext} to {target_ext} not supported'
                }
                
        except Exception as e:
            logger.error(f"Error converting document format: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def is_supported_format(self, file_path: str, input_or_output: str = 'input') -> bool:
        """Check if document format is supported"""
        ext = Path(file_path).suffix.lower()
        supported = self.SUPPORTED_FORMATS.get(input_or_output, [])
        return ext in supported
    
    def get_document_info(self, document_path: str) -> Dict[str, Any]:
        """Get basic document information"""
        try:
            ext = Path(document_path).suffix.lower()
            file_size = os.path.getsize(document_path)
            
            info = {
                'format': ext,
                'file_size_bytes': file_size,
                'file_size_mb': round(file_size / (1024 * 1024), 2)
            }
            
            # Add format-specific info
            if ext == '.pdf':
                try:
                    doc = fitz.open(document_path)
                    info.update({
                        'pages': doc.page_count,
                        'metadata': doc.metadata
                    })
                    doc.close()
                except:
                    pass
            
            elif ext == '.docx':
                try:
                    doc = docx.Document(document_path)
                    info.update({
                        'paragraphs': len(doc.paragraphs),
                        'tables': len(doc.tables),
                        'metadata': {
                            'title': doc.core_properties.title or '',
                            'author': doc.core_properties.author or ''
                        }
                    })
                except:
                    pass
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting document info for {document_path}: {str(e)}")
            return {'error': str(e)}
