import pytest
import asyncio
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock

from services.processing_service.services.document_processor import DocumentProcessor

class TestDocumentProcessor:
    """Test cases for DocumentProcessor"""
    
    @pytest.fixture
    def document_processor(self, temp_dir):
        """Create document processor instance"""
        return DocumentProcessor(str(temp_dir))
    
    @pytest.fixture
    def sample_pdf_file(self, temp_dir):
        """Create a sample PDF file for testing"""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = temp_dir / "sample.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Sample PDF Document")
        c.drawString(100, 730, "This is a test PDF for processing.")
        c.drawString(100, 710, "It contains multiple lines of text.")
        c.drawString(100, 690, "And some test content for extraction.")
        c.showPage()
        c.drawString(100, 750, "Second page content")
        c.drawString(100, 730, "More text on page 2.")
        c.save()
        
        return pdf_path
    
    @pytest.fixture
    def sample_docx_file(self, temp_dir):
        """Create a sample DOCX file for testing"""
        from docx import Document
        
        doc_path = temp_dir / "sample.docx"
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("It contains multiple sentences.")
        doc.add_paragraph("For testing document processing.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        
        doc.save(str(doc_path))
        return doc_path
    
    @pytest.fixture
    def sample_text_file(self, temp_dir):
        """Create a sample text file for testing"""
        text_path = temp_dir / "sample.txt"
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write("This is a test document.\n")
            f.write("It contains multiple lines.\n")
            f.write("For testing text processing.\n")
            f.write("With various content types.\n")
            f.write("Email: test@example.com\n")
            f.write("Phone: (555) 123-4567\n")
            f.write("URL: https://example.com\n")
        
        return text_path
    
    @pytest.fixture
    def sample_excel_file(self, temp_dir):
        """Create a sample Excel file for testing"""
        import openpyxl
        
        excel_path = temp_dir / "sample.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        
        # Add data
        ws['A1'] = "Name"
        ws['B1'] = "Age"
        ws['C1'] = "Email"
        
        ws['A2'] = "John Doe"
        ws['B2'] = 30
        ws['C2'] = "john@example.com"
        
        ws['A3'] = "Jane Smith"
        ws['B3'] = 25
        ws['C3'] = "jane@example.com"
        
        wb.save(str(excel_path))
        return excel_path
    
    @pytest.mark.asyncio
    async def test_extract_text_from_pdf(self, document_processor, sample_pdf_file):
        """Test text extraction from PDF"""
        result = await document_processor.extract_text_from_pdf(str(sample_pdf_file))
        
        assert result['success'] is True
        assert 'text_content' in result
        assert len(result['text_content']) == 2  # 2 pages
        
        # Check extracted content
        first_page = result['text_content'][0]
        assert "Sample PDF Document" in first_page['text']
        assert "test PDF for processing" in first_page['text']
        
        # Check metadata
        assert 'metadata' in result
        assert result['metadata']['page_count'] == 2
    
    @pytest.mark.asyncio
    async def test_extract_text_from_pdf_with_images(self, document_processor, sample_pdf_file):
        """Test text extraction from PDF with image extraction"""
        result = await document_processor.extract_text_from_pdf(
            str(sample_pdf_file),
            extract_images=True
        )
        
        assert result['success'] is True
        assert 'images' in result
        # Note: Our test PDF doesn't have images, so this should be empty
    
    @pytest.mark.asyncio
    async def test_extract_text_from_docx(self, document_processor, sample_docx_file):
        """Test text extraction from DOCX"""
        result = await document_processor.extract_text_from_docx(str(sample_docx_file))
        
        assert result['success'] is True
        assert 'paragraphs' in result
        assert 'tables' in result
        assert 'metadata' in result
        
        # Check paragraphs
        paragraphs = result['paragraphs']
        assert len(paragraphs) >= 4  # Heading + 3 paragraphs
        assert any("Test Document" in para['text'] for para in paragraphs)
        
        # Check tables
        tables = result['tables']
        assert len(tables) == 1
        assert tables[0]['rows'] == 2
        assert tables[0]['columns'] == 3
    
    @pytest.mark.asyncio
    async def test_extract_text_from_excel(self, document_processor, sample_excel_file):
        """Test data extraction from Excel"""
        result = await document_processor.extract_text_from_excel(str(sample_excel_file))
        
        assert result['success'] is True
        assert 'sheets' in result
        assert 'metadata' in result
        
        # Check sheet data
        sheets = result['sheets']
        assert 'Test Sheet' in sheets
        
        sheet_data = sheets['Test Sheet']
        assert sheet_data['rows'] == 3  # Header + 2 data rows
        assert sheet_data['columns'] == 3
        
        # Check data content
        data = sheet_data['data']
        assert data[0] == ['Name', 'Age', 'Email']  # Header
        assert data[1] == ['John Doe', 30, 'john@example.com']
    
    @pytest.mark.asyncio
    async def test_generate_pdf_from_text(self, document_processor, temp_dir, sample_text_file):
        """Test PDF generation from text"""
        output_path = temp_dir / "generated.pdf"
        
        with open(sample_text_file, 'r') as f:
            text_content = f.read()
        
        result = await document_processor.generate_pdf_from_text(
            text_content,
            str(output_path),
            title="Generated Test PDF"
        )
        
        assert result['success'] is True
        assert output_path.exists()
        assert result['title'] == "Generated Test PDF"
        assert result['pages'] > 0
    
    @pytest.mark.asyncio
    async def test_generate_pdf_from_paragraphs(self, document_processor, temp_dir):
        """Test PDF generation from paragraphs"""
        output_path = temp_dir / "generated_from_paragraphs.pdf"
        
        paragraphs = [
            "First paragraph of content.",
            "Second paragraph with more details.",
            "Third paragraph to test formatting."
        ]
        
        result = await document_processor.generate_pdf_from_text(
            paragraphs,
            str(output_path),
            title="Paragraph Test PDF"
        )
        
        assert result['success'] is True
        assert output_path.exists()
        assert result['title'] == "Paragraph Test PDF"
    
    @pytest.mark.asyncio
    async def test_convert_document_format_pdf_to_text(self, document_processor, sample_pdf_file, temp_dir):
        """Test document format conversion from PDF to text"""
        output_path = temp_dir / "converted.txt"
        
        result = await document_processor.convert_document_format(
            str(sample_pdf_file),
            str(output_path),
            "txt"
        )
        
        assert result['success'] is True
        assert output_path.exists()
        assert result['original_format'] == 'PDF'
        assert result['target_format'] == 'TXT'
        
        # Check that text was extracted
        with open(output_path, 'r') as f:
            content = f.read()
        assert "Sample PDF Document" in content
    
    @pytest.mark.asyncio
    async def test_convert_document_format_text_to_pdf(self, document_processor, sample_text_file, temp_dir):
        """Test document format conversion from text to PDF"""
        output_path = temp_dir / "converted.pdf"
        
        result = await document_processor.convert_document_format(
            str(sample_text_file),
            str(output_path),
            "pdf"
        )
        
        assert result['success'] is True
        assert output_path.exists()
        assert result['original_format'] == 'TXT'
        assert result['target_format'] == 'PDF'
    
    @pytest.mark.asyncio
    async def test_convert_document_format_docx_to_pdf(self, document_processor, sample_docx_file, temp_dir):
        """Test document format conversion from DOCX to PDF"""
        output_path = temp_dir / "converted_from_docx.pdf"
        
        result = await document_processor.convert_document_format(
            str(sample_docx_file),
            str(output_path),
            "pdf"
        )
        
        assert result['success'] is True
        assert output_path.exists()
        assert result['original_format'] == 'DOCX'
        assert result['target_format'] == 'PDF'
    
    @pytest.mark.asyncio
    async def test_analyze_text_content(self, document_processor):
        """Test text content analysis"""
        text = """
        This is a sample document for testing. It contains multiple sentences.
        Some words are repeated: testing, testing, testing.
        There are numbers like 123 and 456.
        Email addresses: test@example.com, user@domain.org
        Phone numbers: (555) 123-4567, 555-987-6543
        URLs: https://example.com, http://test.org
        """
        
        result = await document_processor._analyze_text_content(text)
        
        assert result['type'] == 'text'
        assert 'analysis' in result
        
        analysis = result['analysis']
        assert analysis['character_count'] > 0
        assert analysis['word_count'] > 0
        assert analysis['sentence_count'] > 0
        assert analysis['paragraph_count'] > 0
        
        # Check pattern detection
        patterns = analysis['patterns_found']
        assert 'email' in patterns
        assert 'phone' in patterns
        assert 'url' in patterns
        assert 'date' in patterns  # Numbers might be detected as dates
        
        # Check keywords
        keywords = analysis['keywords']
        assert len(keywords) > 0
        assert 'testing' in keywords  # Should be a common word
    
    @pytest.mark.asyncio
    async def test_analyze_text_file(self, document_processor, sample_text_file):
        """Test text file analysis"""
        result = await document_processor.analyze_text_file(str(sample_text_file))
        
        assert result['type'] == 'text'
        assert 'analysis' in result
        
        analysis = result['analysis']
        assert analysis['character_count'] > 0
        assert analysis['word_count'] > 0
        assert analysis['sentence_count'] > 0
    
    def test_is_supported_format_input(self, document_processor):
        """Test supported input format checking"""
        assert document_processor.is_supported_format("test.pdf", "input") is True
        assert document_processor.is_supported_format("test.docx", "input") is True
        assert document_processor.is_supported_format("test.txt", "input") is True
        assert document_processor.is_supported_format("test.xlsx", "input") is True
        assert document_processor.is_supported_format("test.jpg", "input") is False
        assert document_processor.is_supported_format("test.mp4", "input") is False
    
    def test_is_supported_format_output(self, document_processor):
        """Test supported output format checking"""
        assert document_processor.is_supported_format("test.pdf", "output") is True
        assert document_processor.is_supported_format("test.txt", "output") is True
        assert document_processor.is_supported_format("test.docx", "output") is True
        assert document_processor.is_supported_format("test.html", "output") is True
        assert document_processor.is_supported_format("test.jpg", "output") is False
        assert document_processor.is_supported_format("test.xlsx", "output") is False
    
    def test_get_document_info_pdf(self, document_processor, sample_pdf_file):
        """Test getting PDF document info"""
        info = document_processor.get_document_info(str(sample_pdf_file))
        
        assert 'format' in info
        assert 'file_size_bytes' in info
        assert 'file_size_mb' in info
        assert info['format'] == 'pdf'
        assert info['file_size_bytes'] > 0
        assert info['file_size_mb'] > 0
    
    def test_get_document_info_docx(self, document_processor, sample_docx_file):
        """Test getting DOCX document info"""
        info = document_processor.get_document_info(str(sample_docx_file))
        
        assert 'format' in info
        assert 'file_size_bytes' in info
        assert 'file_size_mb' in info
        assert info['format'] == 'docx'
        assert info['file_size_bytes'] > 0
    
    def test_get_document_info_text(self, document_processor, sample_text_file):
        """Test getting text document info"""
        info = document_processor.get_document_info(str(sample_text_file))
        
        assert 'format' in info
        assert 'file_size_bytes' in info
        assert 'file_size_mb' in info
        assert info['format'] == 'txt'
        assert info['file_size_bytes'] > 0
    
    @pytest.mark.asyncio
    async def test_extract_text_from_nonexistent_file(self, document_processor):
        """Test text extraction from non-existent file"""
        result = await document_processor.extract_text_from_pdf("nonexistent.pdf")
        
        assert result['success'] is False
        assert 'error' in result
    
    @pytest.mark.asyncio
    async def test_convert_unsupported_format(self, document_processor, temp_dir):
        """Test conversion of unsupported format"""
        # Create a dummy file with unsupported extension
        dummy_file = temp_dir / "test.xyz"
        dummy_file.write_text("test content")
        
        output_path = temp_dir / "output.pdf"
        
        result = await document_processor.convert_document_format(
            str(dummy_file),
            str(output_path),
            "pdf"
        )
        
        assert result['success'] is False
        assert 'error' in result
        assert "Unsupported" in result['error']
    
    @pytest.mark.asyncio
    async def test_generate_pdf_from_empty_text(self, document_processor, temp_dir):
        """Test PDF generation from empty text"""
        output_path = temp_dir / "empty.pdf"
        
        result = await document_processor.generate_pdf_from_text("", str(output_path))
        
        assert result['success'] is True
        assert output_path.exists()
        assert result['pages'] == 0
    
    @pytest.mark.asyncio
    async def test_extract_text_from_corrupted_file(self, document_processor, temp_dir):
        """Test text extraction from corrupted file"""
        # Create a corrupted PDF file
        corrupted_file = temp_dir / "corrupted.pdf"
        corrupted_file.write_bytes(b"This is not a valid PDF file")
        
        result = await document_processor.extract_text_from_pdf(str(corrupted_file))
        
        assert result['success'] is False
        assert 'error' in result
    
    @pytest.mark.asyncio
    @patch('services.processing_service.services.document_processor.fitz')
    async def test_extract_text_with_fitz_error(self, mock_fitz, document_processor, temp_dir):
        """Test text extraction when PyMuPDF fails"""
        # Mock PyMuPDF to raise an exception
        mock_fitz.open.side_effect = Exception("Mock error")
        
        # Create a dummy PDF file
        pdf_file = temp_dir / "test.pdf"
        pdf_file.write_bytes(b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n>>\nendobj\n")
        
        result = await document_processor.extract_text_from_pdf(str(pdf_file))
        
        assert result['success'] is False
        assert 'error' in result
